from docx import Document

# Load the .docx file
file_path = "IT_Council_1st__1444H_09_02_1443-1.docx"  # Replace with the path to your file
doc = Document(file_path)

# Iterate through paragraphs to extract text
print("Text from the .docx file:")
for paragraph in doc.paragraphs:
    print(paragraph.text)

# If the .docx has tables, extract table data
print("\nTables from the .docx file:")
for table in doc.tables:
    for row in table.rows:
        row_data = [cell.text for cell in row.cells]
        print(row_data)
